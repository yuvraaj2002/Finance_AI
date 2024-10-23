import streamlit as st
from docx import Document
from docx.shared import Pt
from io import BytesIO
import os

st.markdown(
    """
        <style>
               .block-container {
                    padding-top: 1rem;
                    padding-bottom: 0rem;
                }
                .top-margin{
                    margin-top: 4rem;
                    margin-bottom:2rem;
                }
                .block-button{
                    padding: 10px; 
                    width: 100%;
                    background-color: #c4fcce;
                }
        </style>
        """,
    unsafe_allow_html=True,
)


def create_nda_response(
    company_party,
    service_provider_party,
    company_address,
    service_provider_address,
    dates,
    time_period,
    jurisdiction,
    item,
):
    return f"""
# Non-Disclosure Agreement (NDA)

**This Agreement made on this {dates.day} day of {dates.strftime("%B")}, {dates.year}**

By and Between

**{company_party}**, a company incorporated under the Companies Act, 1956 and having its registered office at {company_address} (hereinafter referred to as “Company”, which expression shall unless repugnant to the context or meaning thereof, include its successors in interests and assigns) of the one part;

And

**{service_provider_party}**, a company incorporated under the Companies Act, 2013 and having its registered office at {service_provider_address} (hereinafter referred to as “Service Provider”, which expression shall, unless repugnant to the context or meaning thereof, be deemed to include its representatives and permitted assigns) of the other part;

**Company** and **Service Provider** shall hereinafter be referred to as such or collectively as “Parties” and individually as “Party”.

---

**WHEREAS** both the Parties herein wish to pursue discussions and negotiate with each other for the purpose of entering into a potential business arrangement in relation to {item} (“Proposed Transaction”);

**AND WHEREAS** the Parties contemplate that with respect to the Proposed Transaction, both the Parties may exchange certain information, material and documents relating to each other’s business, assets, financial condition, operations, plans and/or prospects of their businesses (hereinafter referred to as “Confidential Information”, more fully detailed in clause 1 herein below) that each Party regards as proprietary and confidential;

**AND WHEREAS**, each Party wishes to review such Confidential Information of the other for the sole purpose of determining their mutual interest in engaging in the Proposed Transaction;

---

IN CONNECTION WITH THE ABOVE, THE PARTIES HEREBY AGREE AS FOLLOWS:

1.	“Confidential and/or proprietary Information” shall mean and include any information disclosed by one Party (Disclosing Party) to the other (Receiving Party) either directly or indirectly, in writing, orally, by inspection of tangible objects. Confidential information shall include, without limitation, any materials, trade secrets, network information, configurations, trademarks, brand name, know-how, business and marketing plans, financial and operational information, and all other non-public information, material or data relating to the current and/or future business and operations of the Disclosing Party.

2.	The Receiving Party shall refrain from disclosing, reproducing, summarizing and/or distributing Confidential Information and confidential materials of the Disclosing Party except in connection with the Proposed Transaction.

3.	The Parties shall protect the confidentiality of each other’s Confidential Information in the same manner as they protect the confidentiality of their own proprietary and confidential information of similar nature. Each Party, while acknowledging the confidential and proprietary nature of the Confidential Information agrees to take all reasonable measures at its own expense to restrain its representatives from prohibited or unauthorized disclosure or use of the Confidential Information.

4.	Confidential Information shall at all times remain the property of the Disclosing Party and may not be copied or reproduced by the Receiving Party without the Disclosing Party’s prior written consent.

5.	Within seven (7) days of a written request by the Disclosing Party, the Receiving Party shall return/destroy (as may be requested in writing by the Disclosing Party or upon expiry and or earlier termination) all originals, copies, reproductions and summaries of Confidential Information provided to the Receiving Party as Confidential Information. The Receiving Party shall certify to the Disclosing Party in writing that it has satisfied its obligations under this paragraph.

6.	The Receiving Party may disclose the Confidential Information only to the Receiving Party's employees and consultants on a need-to-know basis. The Receiving Party shall have executed or shall execute appropriate written agreements with third parties, in a form and manner sufficient to enable the Receiving Party to enforce all the provisions of this Agreement.

7.	Confidential Information, however, shall not include any information which the Receiving Party can show:
   - i) is in or comes into the public domain otherwise than through a breach of this Agreement or the fault of the Receiving Party; or
   - ii) was already in its possession free of any such restriction prior to receipt from the Disclosing Party; or
   - iii) was independently developed by the Receiving Party without making use of the Confidential Information; or
   - iv) has been approved for release or use (in either case without restriction) by written authorization of the Disclosing Party.

8.	In the event either Party receives a summons or other validly issued administrative or judicial process requiring the disclosure of Confidential Information of the other Party, the Receiving Party shall promptly notify the Disclosing Party. The Receiving Party may disclose Confidential Information to the extent such disclosure is required by law, rule, regulation, or legal process; provided however, that, to the extent practicable, the Receiving Party shall give prompt written notice of any such request for such information to the Disclosing Party, and agrees to cooperate with the Disclosing Party, at the Disclosing Party’s expense, to the extent permissible and practicable, to challenge the request or limit the scope thereof, as the Disclosing Party may reasonably deem appropriate.

9.	Neither Party shall use the other’s name, trademarks, proprietary words or symbols or disclose under this Agreement in any publication, press release, marketing material, or otherwise without the prior written approval of the other.

10.	This Agreement shall be governed by the laws of India. Both parties irrevocably submit to the exclusive jurisdiction of the Courts in {jurisdiction}, for any action or proceeding regarding this Agreement.

---

**IN WITNESS WHEREOF**, the Parties hereto have executed this confidentiality agreement in duplicate by affixing the signature of the authorized representatives as of the date herein above mentioned.

Party 1:		______  
Signature 1:		______  
Name:		______  
Designation:		______  
Place:		______  
Date:		______  

Party 2:		______  
Signature 2:		______  
Name:		______  
Designation:		______  
Place:		______  
Date:		______
"""


def create_docx(nda_response):
    # Create a new Document
    doc = Document()

    # Set the default font style to Times New Roman, 11pt
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    # Add the NDA content
    for paragraph in nda_response.split("\n"):
        if paragraph.startswith("#"):
            # Heading detection based on the number of # symbols
            heading_level = paragraph.count("#")
            heading_text = paragraph.replace("#", "").strip()
            doc.add_heading(heading_text, level=heading_level)
        else:
            # For bold text surrounded by **
            p = doc.add_paragraph()
            if "**" in paragraph:
                parts = paragraph.split("**")
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        p.add_run(part)  # Normal text
                    else:
                        p.add_run(part).bold = True  # Bold text
            else:
                p.add_run(paragraph)  # Add regular text

    # Save the document in a BytesIO object for download
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def nda_creation_form():
    st.title("NDA Creation Form")

    input_tab, output_tab = st.columns(spec=(1, 1), gap="large")
    create_nda = False
    with input_tab:
        st.markdown(
            "<p style='font-size: 18px;'>This module is specifically designed to streamline the process of creating a comprehensive Non-Disclosure Agreement (NDA). It is essential to provide all the necessary information as outlined below to ensure the initiation of the agreement creation process. Please ensure that all details are accurate and complete to facilitate a smooth and efficient process.</p>",
            unsafe_allow_html=True,
        )
        # Input fields for the NDA parameters
        company_party = st.text_input("Company Party")
        service_provider_party = st.text_input("Service Provider Party")
        company_address = st.text_input("Company Address")
        service_provider_address = st.text_input("Service Provider Address")
        dates = st.date_input("Effective Dates")
        time_period = st.text_input("Duration of Agreement")
        jurisdiction = st.text_input("Governing Jurisdiction")
        item = st.text_input("Subject Matter")
        # Submit button
        if st.button("Create NDA", use_container_width=True):
            st.success("NDA details submitted successfully!")
            create_nda = True

    with output_tab:
        if create_nda:
            with st.spinner("Generating response..."):
                nda_response = create_nda_response(
                    company_party,
                    service_provider_party,
                    company_address,
                    service_provider_address,
                    dates,
                    time_period,
                    jurisdiction,
                    item,
                )
                st.markdown(nda_response, unsafe_allow_html=True)

                # Create DOCX file
                buffer = create_docx(nda_response)

                # Provide download button for the generated DOCX
                st.download_button(
                    label="Download NDA as DOCX",
                    data=buffer,
                    file_name="NDA_Agreement.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )


nda_creation_form()
